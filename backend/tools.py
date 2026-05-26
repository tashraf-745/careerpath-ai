import requests
import pypdf
import docx as python_docx
from config import SERPER_API_KEY, RAPIDAPI_KEY, CURRENTS_API_KEY, client, MODEL


def web_search(query: str, num_results: int = 5) -> list:
    url = "https://google.serper.dev/search"
    headers = {"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"}
    payload = {"q": query, "num": num_results}
    response = requests.post(url, headers=headers, json=payload, timeout=10)
    response.raise_for_status()
    return [
        {"title": r.get("title"), "snippet": r.get("snippet"), "link": r.get("link")}
        for r in response.json().get("organic", [])
    ]


def job_search(query: str, location: str = "United States",
               experience_level: str = "entry_level", num_results: int = 10) -> list:
    if not RAPIDAPI_KEY:
        raise RuntimeError("RAPIDAPI_KEY not configured — job_search unavailable")
    url = "https://jsearch.p.rapidapi.com/search"
    headers = {"X-RapidAPI-Key": RAPIDAPI_KEY, "X-RapidAPI-Host": "jsearch.p.rapidapi.com"}
    params = {
        "query":       f"{experience_level.replace('_', ' ')} {query} in {location}",
        "num_pages":   "1",
        "date_posted": "month",
    }
    response = requests.get(url, headers=headers, params=params, timeout=15)
    response.raise_for_status()
    jobs = response.json().get("data", [])[:num_results]
    return [
        {
            "title":       j.get("job_title"),
            "company":     j.get("employer_name"),
            "location":    j.get("job_city"),
            "description": (j.get("job_description") or "")[:500],
            "apply_link":  j.get("job_apply_link"),
            "salary_min":  j.get("job_min_salary"),
            "salary_max":  j.get("job_max_salary"),
            "posted":      j.get("job_posted_at_datetime_utc"),
        }
        for j in jobs
    ]


def get_news(query: str, num_results: int = 20) -> list:
    if not CURRENTS_API_KEY:
        raise RuntimeError("CURRENTS_API_KEY not configured — get_news unavailable")
    url = "https://api.currentsapi.services/v1/search"
    params = {"apiKey": CURRENTS_API_KEY, "keywords": query,
              "language": "en", "page_size": num_results}
    response = requests.get(url, params=params, timeout=10)
    response.raise_for_status()
    return [
        {
            "title":       a.get("title"),
            "description": (a.get("description") or "")[:300],
            "published":   a.get("published"),
            "url":         a.get("url"),
            "category":    a.get("category"),
        }
        for a in response.json().get("news", [])
    ]


def get_salary_benchmark(occupation: str, location: str = "United States") -> dict:
    if RAPIDAPI_KEY:
        try:
            url = "https://jsearch.p.rapidapi.com/search"
            headers = {"X-RapidAPI-Key": RAPIDAPI_KEY, "X-RapidAPI-Host": "jsearch.p.rapidapi.com"}
            params = {"query": f"{occupation} in {location}", "num_pages": "3", "date_posted": "year"}
            response = requests.get(url, headers=headers, params=params, timeout=15)
            response.raise_for_status()
            jobs = response.json().get("data", [])
            salaries = []
            for j in jobs:
                min_s, max_s = j.get("job_min_salary"), j.get("job_max_salary")
                if min_s and max_s:
                    salaries.append((min_s + max_s) / 2)
                elif min_s:
                    salaries.append(min_s)
                elif max_s:
                    salaries.append(max_s)
            if salaries:
                salaries.sort()
                return {
                    "occupation":         occupation,
                    "location":           location,
                    "median_annual":      int(salaries[len(salaries) // 2]),
                    "entry_annual":       int(min(salaries)),
                    "experienced_annual": int(max(salaries)),
                    "sample_size":        len(salaries),
                    "source":             "JSearch live listings",
                }
        except Exception as e:
            print(f"JSearch unavailable ({e}) — falling back to web_search")

    try:
        results = web_search(query=f"{occupation} average salary {location} 2025", num_results=5)
        snippets = " ".join(r.get("snippet", "") for r in results)
        return {
            "occupation":         occupation,
            "location":           location,
            "median_annual":      None,
            "entry_annual":       None,
            "experienced_annual": None,
            "sample_size":        0,
            "source":             "web_search fallback",
            "note":               f"Salary snippets for LLM: {snippets[:600]}",
        }
    except Exception as e:
        return {
            "occupation": occupation, "location": location,
            "median_annual": None, "entry_annual": None, "experienced_annual": None,
            "sample_size": 0, "source": "unavailable",
            "note": f"Both sources failed: {str(e)}",
        }


def parse_resume(file_path: str) -> str:
    file_path = file_path.strip()
    if not file_path:
        return ""
    if file_path.lower().endswith(".pdf"):
        reader = pypdf.PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            return "[PARSE ERROR] Image-based PDF detected. Provide a text-based PDF or DOCX."
        return text
    elif file_path.lower().endswith(".docx"):
        doc = python_docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif file_path.lower().endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: '{file_path}'. Only .pdf, .docx, .txt are supported.")


def parse_resume_bytes(content: bytes, filename: str) -> str:
    import io
    if filename.lower().endswith(".pdf"):
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            return "[PARSE ERROR] Image-based PDF. Please provide a text-based PDF or DOCX."
        return text
    elif filename.lower().endswith(".docx"):
        doc = python_docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    elif filename.lower().endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: '{filename}'")
